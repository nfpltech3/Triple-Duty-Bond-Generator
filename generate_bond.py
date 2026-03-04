"""
Triple Duty Bond Generator - GUI Version
Automatically generates Triple Duty Bond documents from Bill of Entry PDF files.

Features:
- Browse and select Bill of Entry PDF
- Template is embedded/pre-configured (bundled in EXE)
- Choose output location
- One-click generation
- Preserves original formatting (bold, italic, etc.)
"""

import pdfplumber
import re
import math
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import sys
import copy

# ============================================================================
# CONFIGURATION - Template Path (handles both development and EXE modes)
# ============================================================================

def resource_path(relative_path):
    """
    Get absolute path to resource - works for development and PyInstaller EXE.
    When bundled as EXE, PyInstaller extracts files to a temp folder (_MEIPASS).
    """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except AttributeError:
        # Running in development mode
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    return os.path.join(base_path, relative_path)

# Template filename (will be bundled with EXE)
TEMPLATE_FILENAME = "Triple_Duty_Bond_template__IR53112.docx"

def get_default_template_path():
    """Get the default template path - checks bundled location first, then current directory."""
    # First try bundled path (for EXE)
    bundled_path = resource_path(TEMPLATE_FILENAME)
    if os.path.exists(bundled_path):
        return bundled_path
    
    # Try current directory
    current_dir_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), TEMPLATE_FILENAME)
    if os.path.exists(current_dir_path):
        return current_dir_path
    
    # Try hardcoded development path
    dev_path = r"c:\projects\DOCS\Triple duty bond\Triple_Duty_Bond_template__IR53112.docx"
    if os.path.exists(dev_path):
        return dev_path
    
    return ""  # Not found

DEFAULT_TEMPLATE_PATH = get_default_template_path()


class BillOfEntryExtractor:
    """Extracts relevant data from Bill of Entry PDF files."""
    
    def __init__(self, pdf_path):
        self.pdf_path = pdf_path
        self.full_text = ""
        self.data = {}
        
    def load_pdf(self):
        """Load and extract text from PDF."""
        with pdfplumber.open(self.pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    self.full_text += text + "\n"
        return self
    
    def extract_all(self):
        """Extract all relevant fields from the PDF."""
        self.load_pdf()
        
        self.data = {
            'be_no': self._extract_be_no(),
            'be_date': self._extract_be_date(),
            'be_date_formatted': self._format_be_date(),
            'importer_name': self._extract_importer_name(),
            'iec_no': self._extract_iec_no(),
            'port_code': 'INBOM4',
            'port_name': 'MUMBAI Sahar Air Cargo',
            'wh_code': self._extract_wh_code(),
            'debt_amt': self._extract_debt_amt(),
            'total_assessed_value': self._extract_total_assessed_value(),
            'total_packages': self._extract_packages(),
            'total_packages_num': self._extract_packages_num(),
            'invoice_count': self._extract_invoice_count(),
        }
        return self.data
    
    def _extract_be_no(self):
        match = re.search(r'INBOM4\s+(\d+)\s+\d{2}/\d{2}/\d{4}', self.full_text)
        return match.group(1) if match else ''
    
    def _extract_be_date(self):
        match = re.search(r'INBOM4\s+\d+\s+(\d{2}/\d{2}/\d{4})', self.full_text)
        return match.group(1) if match else ''
    
    def _format_be_date(self):
        date_str = self._extract_be_date()
        if date_str:
            try:
                date_obj = datetime.strptime(date_str, '%d/%m/%Y')
                return date_obj.strftime('%d-%b-%Y')
            except:
                pass
        return date_str
    
    def _extract_importer_name(self):
        match = re.search(r'1\.IMPORTER NAME.*?\n([A-Z][A-Z\s]+(?:PRIVATE|PVT)?\.?\s*(?:LIMITED|LTD)?\.?)', self.full_text)
        if match:
            name = match.group(1).strip()
            name = re.sub(r'\s+[A-Z]$', '', name)
            return name
        return ''
    
    def _extract_iec_no(self):
        match = re.search(r'IEC/Br\s+(\d+)', self.full_text)
        return match.group(1) if match else ''
    
    def _extract_wh_code(self):
        match = re.search(r'([EN]SA\d[A-Z]\d{3})', self.full_text)
        return match.group(1) if match else ''
    
    def _extract_debt_amt(self):
        """Extract DEBT AMT - already 3x the total duty."""
        match = re.search(r'INBOM4\s+WH\s+(\d+)', self.full_text)
        return int(match.group(1)) if match else 0
    
    def _extract_total_assessed_value(self):
        """
        Extract TOT.ASS VAL (18.TOT.ASS VAL) from the Bill of Entry.
        This is the total assessed value that appears in the Value Rs. column.
        
        The value appears on the line with 1.BCD, 2.ACD... 18.TOT.ASS VAL
        It's the last number on the next line (can be 6-8 digits, may have decimal)
        """
        # Pattern: Find the line with 8.G.CESS and 18.TOT.ASS VAL, then get values from next line
        # The line format is: value1 value2 value3 value4 value5 value6 value7 value8
        # Where value8 (last) is TOT.ASS VAL
        
        # Look for the duty summary line - 8 values, last one is TOT.ASS VAL
        # Pattern: 8.G.CESS followed by 18.TOT.ASS VAL, then a line of numbers
        match = re.search(
            r'8\.G\.CESS\s+18\.TOT\.ASS VAL\s*\n'
            r'([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)',
            self.full_text
        )
        if match:
            # Last value (group 8) is TOT.ASS VAL
            val = match.group(8)
            return int(float(val))  # Handle decimal values like 524174.49
        
        # Alternative pattern: Look for 8 numbers after the header line
        match = re.search(
            r'18\.TOT\.ASS VAL\s*\n'
            r'([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)\s+([\d.]+)',
            self.full_text
        )
        if match:
            val = match.group(8)
            return int(float(val))
        
        # Fallback: Find any large number (6-8 digits) that could be the value
        # after 18.TOT.ASS VAL
        match = re.search(r'18\.TOT\.ASS VAL.*?\n.*?(\d{6,8})(?:\.\d+)?', self.full_text, re.DOTALL)
        if match:
            return int(match.group(1))
        
        return 0
    
    def _extract_invoice_count(self):
        """
        Extract the number of invoices from the Bill of Entry.
        This appears in the TYPE INV ITEM CONT line.
        """
        # Pattern: TYPE INV ITEM CONT followed by Nos X Y Z where X is invoice count
        # or Nos followed by 3 numbers
        match = re.search(r'TYPE\s+INV\s+ITEM\s+CONT.*?Nos\s+(\d+)', self.full_text, re.DOTALL)
        if match:
            return match.group(1)
        
        # Alternative: Look for "Nos X Y Z" pattern
        match = re.search(r'Nos\s+(\d+)\s+\d+\s+\d+', self.full_text)
        if match:
            return match.group(1)
        
        return '1'  # Default to 1 invoice
    
    def _extract_packages(self):
        match = re.search(r'BE PKG\s+(\d+)', self.full_text)
        return match.group(1) + ' PKG' if match else ''
    
    def _extract_packages_num(self):
        match = re.search(r'BE PKG\s+(\d+)', self.full_text)
        return int(match.group(1)) if match else 0


class BondCalculator:
    """Handles bond amount calculations."""
    
    @staticmethod
    def calculate_bond_amount(debt_amt, round_to=5000):
        """Round up DEBT AMT to nearest round_to value (default: 5,000)."""
        return math.ceil(debt_amt / round_to) * round_to
    
    @staticmethod
    def format_indian_amount(amount):
        """Format amount in Indian numbering system."""
        amount_str = str(amount)
        if len(amount_str) > 3:
            last_three = amount_str[-3:]
            remaining = amount_str[:-3]
            formatted_remaining = ''
            for i, digit in enumerate(reversed(remaining)):
                if i > 0 and i % 2 == 0:
                    formatted_remaining = ',' + formatted_remaining
                formatted_remaining = digit + formatted_remaining
            formatted = formatted_remaining + ',' + last_three
        else:
            formatted = amount_str
        return f"Rs.{formatted}/-"
    
    @staticmethod
    def amount_to_words(amount):
        """Convert amount to words in Indian numbering system."""
        ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine',
                'Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen',
                'Seventeen', 'Eighteen', 'Nineteen']
        tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
        
        if amount == 0:
            return 'Zero'
        
        def two_digits(n):
            if n < 20:
                return ones[n]
            return tens[n // 10] + ('' if n % 10 == 0 else '-' + ones[n % 10])
        
        def three_digits(n):
            if n < 100:
                return two_digits(n)
            return ones[n // 100] + ' Hundred' + ('' if n % 100 == 0 else ' ' + two_digits(n % 100))
        
        crore = amount // 10000000
        amount %= 10000000
        lakh = amount // 100000
        amount %= 100000
        thousand = amount // 1000
        amount %= 1000
        remainder = amount
        
        result = []
        if crore:
            result.append(three_digits(crore) + ' Crore')
        if lakh:
            result.append(three_digits(lakh) + ' Lakh')
        if thousand:
            result.append(three_digits(thousand) + ' Thousand')
        if remainder:
            result.append(three_digits(remainder))
        
        return ' '.join(result)


class TripleDutyBondGenerator:
    """Generates Triple Duty Bond documents from templates."""
    
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = None
        
    def load_template(self):
        """Load the Word template."""
        self.doc = Document(self.template_path)
        return self
    
    def replace_in_runs_preserve_format(self, paragraph, old_text, new_text):
        """
        Replace text in paragraph while PRESERVING formatting.
        Handles text split across multiple runs.
        """
        # Get full paragraph text
        full_text = ''.join(run.text for run in paragraph.runs)
        
        if old_text not in full_text:
            return False
        
        # Find where replacement needs to happen
        start_idx = full_text.find(old_text)
        end_idx = start_idx + len(old_text)
        
        # Build character-to-run mapping
        char_to_run = []
        for run_idx, run in enumerate(paragraph.runs):
            for _ in run.text:
                char_to_run.append(run_idx)
        
        if not char_to_run:
            return False
        
        # Find which runs contain the old text
        start_run_idx = char_to_run[start_idx]
        end_run_idx = char_to_run[end_idx - 1]
        
        # Calculate position within first run
        chars_before_start_run = sum(len(paragraph.runs[i].text) for i in range(start_run_idx))
        pos_in_start_run = start_idx - chars_before_start_run
        
        # Calculate position within last run
        chars_before_end_run = sum(len(paragraph.runs[i].text) for i in range(end_run_idx))
        pos_in_end_run = end_idx - chars_before_end_run
        
        if start_run_idx == end_run_idx:
            # Old text is entirely within one run - simple replacement
            run = paragraph.runs[start_run_idx]
            run.text = run.text[:pos_in_start_run] + new_text + run.text[pos_in_end_run:]
        else:
            # Old text spans multiple runs
            # Put replacement text in first run, clear others
            first_run = paragraph.runs[start_run_idx]
            first_run.text = first_run.text[:pos_in_start_run] + new_text
            
            # Clear middle runs
            for i in range(start_run_idx + 1, end_run_idx):
                paragraph.runs[i].text = ""
            
            # Trim last run
            last_run = paragraph.runs[end_run_idx]
            last_run.text = last_run.text[pos_in_end_run:]
        
        return True
    
    def replace_text_preserve_format(self, old_text, new_text):
        """Replace text everywhere while preserving formatting."""
        replaced = False
        
        # Process paragraphs
        for para in self.doc.paragraphs:
            if old_text in para.text:
                if self.replace_in_runs_preserve_format(para, old_text, new_text):
                    replaced = True
        
        # Process tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            if self.replace_in_runs_preserve_format(para, old_text, new_text):
                                replaced = True
        
        return replaced
    
    def _replace_table_cell_value(self, row_idx, cell_idx, new_value):
        """
        Replace the content of a specific table cell.
        This is used for short values like invoice count that would
        cause issues with global text replacement.
        """
        for table in self.doc.tables:
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if cell_idx < len(row.cells):
                    cell = row.cells[cell_idx]
                    # Clear existing content and set new value
                    for para in cell.paragraphs:
                        if para.runs:
                            # Keep first run's formatting, clear text
                            first_run = para.runs[0]
                            for run in para.runs:
                                run.text = ""
                            first_run.text = str(new_value)
                        else:
                            # No runs, just set the text
                            para.text = str(new_value)
    
    def generate(self, be_data, output_path):
        """Generate the Triple Duty Bond document."""
        self.load_template()
        
        # Calculate bond amount
        calc = BondCalculator()
        bond_amount = calc.calculate_bond_amount(be_data['debt_amt'])
        bond_formatted = calc.format_indian_amount(bond_amount)
        bond_words = calc.amount_to_words(bond_amount)
        
        # Parse dates
        date_parts = be_data['be_date_formatted'].split('-')
        day, month, year = date_parts[0], date_parts[1], date_parts[2]
        day_int = int(day)
        
        if day_int in [1, 21, 31]:
            day_suffix = "st"
        elif day_int in [2, 22]:
            day_suffix = "nd"
        elif day_int in [3, 23]:
            day_suffix = "rd"
        else:
            day_suffix = "th"
        
        # Define all replacements (template values -> new values)
        # Order matters - do longer strings first to avoid partial replacements
        # Note: Template uses \xa0 (non-breaking space) in some places
        replacements = [
            # Bond amount with words (longer string first)
            ('Rs. Twenty-Four Lakh Fifty Thousand Only', f'Rs. {bond_words} Only'),
            
            # Bond amount formatted
            ('Rs.24,50,000/-', bond_formatted),
            
            # "Sealed with" date line - uses non-breaking spaces (\xa0)
            # Template: 'Sealed\xa0with\xa0our\xa0seal(s)\xa0this  03rd   day\xa0of Feb-2026.'
            (' 03rd   day\xa0of Feb-2026', f" {day}{day_suffix}   day\xa0of {month}-{year}"),
            
            # BE reference date (in "Bill of Entry No. xxx dt. dd-Mon-yyyy")
            ('03-Feb-2026', be_data['be_date_formatted']),
            
            # BE Number
            ('7281285', be_data['be_no']),
            
            # Value Rs. in table (TOT.ASS VAL)
            ('1855456', str(be_data['total_assessed_value'])),
            
            # Sl No. of invoice in table (invoice count)
            # Note: This is in a table cell, value is just "1"
            # We need to be careful to only replace the exact cell value
        ]
        
        # Apply all replacements (preserving formatting)
        for old_text, new_text in replacements:
            self.replace_text_preserve_format(old_text, new_text)
        
        # Handle invoice count and packages separately for table cells
        # These need special handling as "1" and "2" are very short strings
        self._replace_table_cell_value(2, 2, be_data['invoice_count'])  # Sl No of invoice
        self._replace_table_cell_value(2, 3, be_data['total_packages'])  # Packages
        
        # Save document
        self.doc.save(output_path)
        
        return {
            'output_path': output_path,
            'be_no': be_data['be_no'],
            'be_date': be_data['be_date_formatted'],
            'bond_amount': bond_amount,
            'bond_formatted': bond_formatted,
            'bond_words': bond_words,
            'total_assessed_value': be_data['total_assessed_value'],
            'packages': be_data['total_packages'],
        }


class TripleDutyBondGUI:
    """GUI Application for Triple Duty Bond Generator."""
    
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Triple Duty Bond Generator")
        self.root.geometry("900x700")
        try:
            self.root.state('zoomed')  # Full screen on Windows
        except:
            pass
        self.root.configure(bg='#F4F6F8')
        
        # Configure style - Nagarkot Brand Colors
        self.style = ttk.Style()
        if 'clam' in self.style.theme_names():
            self.style.theme_use('clam')
            
        self.style.configure('TFrame', background='#F4F6F8')
        self.style.configure('Panel.TFrame', background='#FFFFFF')
        
        self.style.configure('Title.TLabel', font=('Segoe UI', 22, 'bold'), foreground='#1E1E1E', background='#FFFFFF')
        self.style.configure('Subtitle.TLabel', font=('Segoe UI', 11), foreground='#6B7280', background='#FFFFFF')
        self.style.configure('Heading.TLabel', font=('Segoe UI', 10, 'bold'), foreground='#1F3F6E', background='#FFFFFF')
        self.style.configure('Info.TLabel', font=('Segoe UI', 9), foreground='#6B7280', background='#FFFFFF')
        self.style.configure('Success.TLabel', font=('Segoe UI', 9, 'bold'), foreground='#1F3F6E', background='#FFFFFF')
        self.style.configure('Footer.TLabel', font=('Segoe UI', 8), foreground='#6B7280', background='#F4F6F8')
        self.style.configure('TLabel', background='#FFFFFF', foreground='#1E1E1E')
        
        self.style.configure('Generate.TButton', font=('Segoe UI', 11, 'bold'), 
                             background='#1F3F6E', foreground='#FFFFFF', padding=10)
        self.style.map('Generate.TButton', background=[('active', '#2A528F')])
        
        self.style.configure('TButton', font=('Segoe UI', 9), background='#FFFFFF', foreground='#1F3F6E')
        
        self.style.configure('TLabelframe', background='#FFFFFF', bordercolor='#E5E7EB')
        self.style.configure('TLabelframe.Label', background='#FFFFFF', foreground='#1F3F6E', font=('Segoe UI', 10, 'bold'))
        self.style.configure('Panel.TCheckbutton', background='#FFFFFF')
        
        # Variables
        self.pdf_path = tk.StringVar()
        self.template_path = tk.StringVar(value=DEFAULT_TEMPLATE_PATH)
        self.output_path = tk.StringVar()
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface adhering to Nagarkot Brand Standard."""
        # 1. HEADER (Dynamic Height, Panel White)
        header_frame = tk.Frame(self.root, bg='#FFFFFF', height=80)
        header_frame.pack(fill=tk.X, side=tk.TOP)
        header_frame.pack_propagate(False)
        
        # Logo placeholder (Left)
        try:
            from PIL import Image, ImageTk
            logo_path = resource_path("logo.png")
            if os.path.exists(logo_path):
                img = Image.open(logo_path)
                aspect_ratio = img.width / img.height
                new_height = 30
                new_width = int(new_height * aspect_ratio)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                self.logo_img = ImageTk.PhotoImage(img)
                logo_lbl = tk.Label(header_frame, image=self.logo_img, bg='#FFFFFF')
            else:
                logo_lbl = tk.Label(header_frame, text="NAGARKOT", font=('Segoe UI', 16, 'bold', 'italic'), bg='#FFFFFF', fg='#D8232A')
        except:
            logo_lbl = tk.Label(header_frame, text="NAGARKOT", font=('Segoe UI', 16, 'bold', 'italic'), bg='#FFFFFF', fg='#D8232A')
            
        logo_lbl.place(x=30, rely=0.5, anchor=tk.W)
        
        # Centered Title & Subtitle
        title_lbl = tk.Label(header_frame, text="TRIPLE DUTY BOND GENERATOR", font=('Segoe UI', 16, 'bold'), bg='#FFFFFF', fg='#1F3F6E')
        title_lbl.place(relx=0.5, rely=0.35, anchor=tk.CENTER)
        
        subtitle_lbl = tk.Label(header_frame, text="Generate Documents from Bill of Entry PDFs", font=('Segoe UI', 10), bg='#FFFFFF', fg='#6B7280')
        subtitle_lbl.place(relx=0.5, rely=0.7, anchor=tk.CENTER)
        
        # Divider Line
        tk.Frame(self.root, bg='#E5E7EB', height=1).pack(fill=tk.X)
        
        # 3. FOOTER (Bottom-Left)
        footer_frame = ttk.Frame(self.root, style='TFrame')
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM, padx=20, pady=10)
        footer_lbl = ttk.Label(footer_frame, text="Nagarkot Forwarders Pvt. Ltd. ©", style='Footer.TLabel')
        footer_lbl.pack(side=tk.LEFT)
        
        # 2. BODY (Flexible Content Area)
        body_container = ttk.Frame(self.root, style='TFrame')
        body_container.pack(fill=tk.BOTH, expand=True, padx=40, pady=40)
        
        # Main Panel
        main_panel = ttk.Frame(body_container, style='Panel.TFrame')
        main_panel.place(relx=0.5, rely=0.5, anchor=tk.CENTER, relwidth=0.8, relheight=0.9)
        
        # Inner padding frame for the panel
        inner_frame = ttk.Frame(main_panel, padding="30", style='Panel.TFrame')
        inner_frame.pack(fill=tk.BOTH, expand=True)
        
        # Status Label
        if os.path.exists(DEFAULT_TEMPLATE_PATH):
            ttk.Label(inner_frame, text="✓ Template loaded successfully", style='Success.TLabel').pack(pady=(0, 15))
        else:
            ttk.Label(inner_frame, text="⚠ Template not found - please browse to select", style='Info.TLabel').pack(pady=(0, 15))
            
        # File Selection Frame
        file_frame = ttk.LabelFrame(inner_frame, text=" File Configuration ", padding="20")
        file_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Grid layout for inputs
        file_frame.columnconfigure(1, weight=1)
        
        # Bill of Entry PDF
        ttk.Label(file_frame, text="Bill of Entry PDF:", style='Heading.TLabel').grid(row=0, column=0, sticky='w', pady=10, padx=(0,10))
        ttk.Entry(file_frame, textvariable=self.pdf_path).grid(row=0, column=1, sticky='ew', pady=10)
        ttk.Button(file_frame, text="Browse...", command=self.browse_pdf).grid(row=0, column=2, padx=(10,0), pady=10)
        
        # Output Location
        ttk.Label(file_frame, text="Save Output To:", style='Heading.TLabel').grid(row=1, column=0, sticky='w', pady=10, padx=(0,10))
        ttk.Entry(file_frame, textvariable=self.output_path).grid(row=1, column=1, sticky='ew', pady=10)
        ttk.Button(file_frame, text="Browse...", command=self.browse_output).grid(row=1, column=2, padx=(10,0), pady=10)
        
        # Template selection
        self.show_template = tk.BooleanVar(value=False)
        template_check = ttk.Checkbutton(file_frame, text="Change template file", 
                                          variable=self.show_template, command=self.toggle_template, style='Panel.TCheckbutton')
        template_check.grid(row=2, column=0, columnspan=3, sticky='w', pady=5)
        
        self.template_frame = ttk.Frame(file_frame, style='Panel.TFrame')
        self.template_frame.grid(row=3, column=0, columnspan=3, sticky='ew')
        self.template_frame.columnconfigure(1, weight=1)
        
        ttk.Label(self.template_frame, text="Word Template:", style='Heading.TLabel').grid(row=0, column=0, sticky='w', pady=10, padx=(0,10))
        ttk.Entry(self.template_frame, textvariable=self.template_path).grid(row=0, column=1, sticky='ew', pady=10)
        ttk.Button(self.template_frame, text="Browse...", command=self.browse_template).grid(row=0, column=2, padx=(10,0), pady=10)
        self.template_frame.grid_remove()
        
        # Generate Button
        generate_btn = ttk.Button(inner_frame, text="⚙ Generate Triple Duty Bond", 
                                  style='Generate.TButton', command=self.generate_bond)
        generate_btn.pack(pady=(0, 20))
        
        # Status/Log Area
        log_frame = ttk.LabelFrame(inner_frame, text=" Process Logs ", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.status_text = tk.Text(log_frame, height=8, wrap=tk.WORD, font=('Consolas', 9), 
                                   bg='#F9FAFB', fg='#1E1E1E', relief=tk.FLAT, borderwidth=1, highlightthickness=1, highlightcolor='#E5E7EB')
        scrollbar = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.status_text.yview)
        self.status_text.configure(yscrollcommand=scrollbar.set)
        
        self.status_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Initial message
        self.log("Ready. Select the Bill of Entry PDF and click Generate.")
        self.log(f"Template: {os.path.basename(DEFAULT_TEMPLATE_PATH)}")
        self.log("✓ Formatting will be preserved (bold, italic, etc.)")
        
    def toggle_template(self):
        """Show/hide template selection."""
        if self.show_template.get():
            self.template_frame.grid()
        else:
            self.template_frame.grid_remove()
        
    def log(self, message):
        """Log a message to the status text area."""
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)
        self.root.update_idletasks()
        
    def clear_log(self):
        """Clear the status text area."""
        self.status_text.delete(1.0, tk.END)
        
    def browse_pdf(self):
        """Browse for Bill of Entry PDF file."""
        filename = filedialog.askopenfilename(
            title="Select Bill of Entry PDF",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
        )
        if filename:
            self.pdf_path.set(filename)
            # Extract BE number from PDF for output filename
            try:
                import pdfplumber
                with pdfplumber.open(filename) as pdf:
                    text = pdf.pages[0].extract_text() or ""
                    be_match = re.search(r'INBOM4\s+(\d+)\s+\d{2}/\d{2}/\d{4}', text)
                    if be_match:
                        be_no = be_match.group(1)
                        output_name = f"Triple_Duty_Bond_BE_{be_no}.docx"
                    else:
                        base_name = os.path.splitext(os.path.basename(filename))[0]
                        output_name = f"Triple_Duty_Bond_{base_name}.docx"
            except:
                base_name = os.path.splitext(os.path.basename(filename))[0]
                output_name = f"Triple_Duty_Bond_{base_name}.docx"
            output_path = os.path.join(os.path.dirname(filename), output_name)
            self.output_path.set(output_path)
            
    def browse_template(self):
        """Browse for template file."""
        filename = filedialog.askopenfilename(
            title="Select Template File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.template_path.set(filename)
            
    def browse_output(self):
        """Browse for output file location."""
        filename = filedialog.asksaveasfilename(
            title="Save Output As",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.output_path.set(filename)
            
    def generate_bond(self):
        """Generate the Triple Duty Bond document."""
        # Validate inputs
        if not self.pdf_path.get():
            messagebox.showerror("Error", "Please select a Bill of Entry PDF file.")
            return
        if not self.template_path.get():
            messagebox.showerror("Error", "Please select a template file.")
            return
        if not self.output_path.get():
            messagebox.showerror("Error", "Please specify an output file location.")
            return
            
        if not os.path.exists(self.pdf_path.get()):
            messagebox.showerror("Error", f"PDF file not found:\n{self.pdf_path.get()}")
            return
        if not os.path.exists(self.template_path.get()):
            messagebox.showerror("Error", f"Template file not found:\n{self.template_path.get()}")
            return
        
        self.clear_log()
        self.log("=" * 60)
        self.log("TRIPLE DUTY BOND GENERATOR")
        self.log("=" * 60)
        
        try:
            # Extract data from PDF
            self.log(f"\n[1] Extracting data from PDF...")
            self.log(f"    File: {os.path.basename(self.pdf_path.get())}")
            
            extractor = BillOfEntryExtractor(self.pdf_path.get())
            be_data = extractor.extract_all()
            
            self.log("\n    Extracted Fields:")
            self.log(f"      BE Number              : {be_data['be_no']}")
            self.log(f"      BE Date                : {be_data['be_date_formatted']}")
            self.log(f"      Importer               : {be_data['importer_name']}")
            self.log(f"      IEC No                 : {be_data['iec_no']}")
            self.log(f"      Port                   : {be_data['port_name']}")
            self.log(f"      WH Code                : {be_data['wh_code']}")
            self.log(f"      Debt Amount            : {BondCalculator.format_indian_amount(be_data['debt_amt'])}")
            self.log(f"      TOT.ASS VAL (Value Rs.): {be_data['total_assessed_value']}")
            self.log(f"      Packages               : {be_data['total_packages']}")
            
            # Calculate bond amount
            self.log("\n[2] Calculating Bond Amount:")
            calc = BondCalculator()
            bond_amount = calc.calculate_bond_amount(be_data['debt_amt'])
            
            self.log(f"      Debt Amount (from BE)   : {calc.format_indian_amount(be_data['debt_amt'])}")
            self.log(f"      (Already 3x - just round up)")
            self.log(f"      Rounded Bond Amount     : {calc.format_indian_amount(bond_amount)}")
            self.log(f"      In Words                : Rs. {calc.amount_to_words(bond_amount)} Only")
            
            # Generate document
            self.log("\n[3] Generating document (preserving formatting)...")
            self.log(f"    Template: {os.path.basename(self.template_path.get())}")
            
            generator = TripleDutyBondGenerator(self.template_path.get())
            result = generator.generate(be_data, self.output_path.get())
            
            self.log("\n    Replacements applied:")
            self.log(f"      Bond Amount  : Rs.24,50,000/- → {result['bond_formatted']}")
            self.log(f"      BE Number    : 7281285 → {result['be_no']}")
            self.log(f"      BE Date      : 03-Feb-2026 → {result['be_date']}")
            self.log(f"      Sealed Date  : 03rd day of Feb-2026 → {result['be_date']}")
            self.log(f"      Value Rs.    : 1855456 → {result['total_assessed_value']}")
            self.log(f"      Packages     : 2 PKG → {result['packages']}")
            
            self.log("\n" + "=" * 60)
            self.log("✅ DOCUMENT GENERATED SUCCESSFULLY!")
            self.log("=" * 60)
            self.log(f"\n    Output File: {os.path.basename(result['output_path'])}")
            self.log(f"    BE Number: {result['be_no']}")
            self.log(f"    BE Date: {result['be_date']}")
            self.log(f"    Bond Amount: {result['bond_formatted']}")
            self.log(f"    Value Rs.: {result['total_assessed_value']}")
            self.log(f"    Packages: {result['packages']}")
            self.log(f"\n    ✓ Bold formatting preserved")
            self.log(f"    The document is ready for use!")
            
            if messagebox.askyesno("Success", 
                f"Document generated successfully!\n\n"
                f"BE Number: {result['be_no']}\n"
                f"BE Date: {result['be_date']}\n"
                f"Bond Amount: {result['bond_formatted']}\n"
                f"Value Rs.: {result['total_assessed_value']}\n"
                f"Packages: {result['packages']}\n\n"
                f"Formatting preserved ✓\n\n"
                f"Do you want to open the output folder?"):
                os.startfile(os.path.dirname(self.output_path.get()))
                
        except Exception as e:
            self.log(f"\n❌ ERROR: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Error", f"An error occurred:\n\n{str(e)}")
    
    def run(self):
        """Start the application."""
        self.root.mainloop()


if __name__ == "__main__":
    app = TripleDutyBondGUI()
    app.run()
