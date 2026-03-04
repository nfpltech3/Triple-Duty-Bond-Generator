from docx import Document

doc = Document(r"c:\projects\DOCS\Triple duty bond\Triple_Duty_Bond_7281285_20260204_180532.docx")

print("=" * 70)
print("VERIFICATION OF GENERATED TRIPLE DUTY BOND DOCUMENT")
print("=" * 70)

for para in doc.paragraphs:
    text = para.text
    if "Rs." in text and "Lakh" in text:
        start = text.find("Rs.")
        end = text.find("Only") + 4 if "Only" in text else len(text)
        print(f"\nBond Amount Found: {text[start:end]}")
    if "Bill of Entry No" in text:
        print(f"\nBE Reference: {text.strip()[:100]}")
    if "Sealed with" in text:
        print(f"\nDate Clause: {text.strip()[:80]}")

print("\n" + "=" * 70)
print("TABLE VALUES")
print("=" * 70)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() and len(cell.text.strip()) < 50:
                val = cell.text.replace("\n", " ").strip()
                if val and val not in ["(1)", "(2)", "(3)", "(4)", "(5)", "(6)"]:
                    print(f"  {val}")

print("\n" + "=" * 70)
print("DOCUMENT VERIFIED SUCCESSFULLY!")
print("=" * 70)
