from docx import Document

doc = Document('Triple_Duty_Bond_BE_5445010.docx')

print("=== Generated Document Content ===")
print()

# Check table
print("TABLE CONTENT:")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.replace('\n', ' | ').strip()
            if text:
                print(f"  Cell: {text}")
        print()

print()
print("KEY PARAGRAPHS:")
for para in doc.paragraphs:
    text = para.text.strip()
    if 'Rs.' in text and 'Lakh' in text:
        print(f"  Bond Amount: {text[:150]}...")
    if 'Bill of Entry' in text:
        print(f"  BE Reference: {text}")
    if 'Sealed' in text:
        print(f"  Sealed: {text}")

# Also check template for comparison
print()
print("="*70)
print("TEMPLATE VALUES (for reference):")
template = Document('Triple_Duty_Bond_template__IR53112.docx')
for table in template.tables:
    print("Template Table:")
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.replace('\n', ' | ').strip()
            if text:
                print(f"  Cell: {text}")
